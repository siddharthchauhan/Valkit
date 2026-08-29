"""Rendering validation documents.

The generator is deliberately unforgiving. Jinja runs with ``StrictUndefined``,
so a template referring to a variable the context does not supply raises rather
than rendering an empty string. In ordinary web templating an undefined
variable producing nothing is a convenience; in a document that will be
electronically signed it is a defect that nobody sees.

Rendering is deterministic: the same records and the same clock produce
byte-identical output, which is what allows a regenerated document to be
compared against the signed one to prove nothing has changed.
"""

from __future__ import annotations

import html as html_module
import re
from pathlib import Path
from typing import Any, Sequence

from jinja2 import Environment, FileSystemLoader, StrictUndefined, TemplateNotFound
from jinja2 import select_autoescape  # noqa: F401 - imported for discoverability

from ..errors import DocumentError
from ..models import (
    AgentSpec,
    Document,
    DocumentStatus,
    DocumentType,
)
from ..util import Clock, SystemClock, sha256_text
from .context import DocumentContext, build_context
from .filters import FILTERS

__all__ = ["DocumentGenerator", "TEMPLATE_DIR", "PACKAGE_ORDER", "markdown_to_html"]

TEMPLATE_DIR = Path(__file__).parent.parent / "templates"

# Generation order. Later documents cite earlier ones, and a reviewer reads
# them in this order, so the package is produced in it too.
PACKAGE_ORDER: tuple[DocumentType, ...] = (
    DocumentType.URS,
    DocumentType.FRS,
    DocumentType.RISK_ASSESSMENT,
    DocumentType.VALIDATION_PLAN,
    DocumentType.CREDIBILITY_PLAN,
    DocumentType.IQ_PROTOCOL,
    DocumentType.OQ_PROTOCOL,
    DocumentType.PQ_PROTOCOL,
    DocumentType.IQ_REPORT,
    DocumentType.OQ_REPORT,
    DocumentType.PQ_REPORT,
    DocumentType.CREDIBILITY_REPORT,
    DocumentType.RTM,
    DocumentType.VSR,
)


class DocumentGenerator:
    """Renders :class:`~valkit.models.Document` records from templates."""

    def __init__(
        self,
        clock: Clock | None = None,
        template_dir: str | Path | None = None,
        *,
        vault: Any = None,
        audit: Any = None,
    ):
        self._clock = clock or SystemClock()
        self._template_dir = Path(template_dir) if template_dir else TEMPLATE_DIR
        self._vault = vault
        self._audit = audit
        self._counter = 0

        self.environment = Environment(
            loader=FileSystemLoader(str(self._template_dir)),
            undefined=StrictUndefined,
            trim_blocks=True,
            lstrip_blocks=True,
            keep_trailing_newline=True,
            # Markdown, not HTML: escaping would corrupt the output.
            autoescape=False,
        )
        self.environment.filters.update(FILTERS)

    # -- generation --------------------------------------------------------

    def template_name(self, doc_type: DocumentType) -> str:
        return f"{doc_type.value.lower()}.md.j2"

    def generate(
        self,
        doc_type: DocumentType,
        context: DocumentContext,
        *,
        doc_id: str | None = None,
        version: str = "1.0",
    ) -> Document:
        """Render one document."""
        self._counter += 1
        doc_id = doc_id or f"DOC-{doc_type.value}-{self._counter:04d}"
        context.doc_id = doc_id
        context.doc_type = doc_type
        context.version = version
        context.document_title = _title(doc_type, context.spec)

        template_name = self.template_name(doc_type)
        try:
            template = self.environment.get_template(template_name)
        except TemplateNotFound as error:
            raise DocumentError(
                f"no template for document type {doc_type.value} (expected "
                f"{template_name} in {self._template_dir})"
            ) from error

        try:
            content = template.render(**context.to_dict())
        except Exception as error:
            raise DocumentError(
                f"could not render {doc_type.value} ({template_name}): {error}"
            ) from error

        content = _normalise(content)
        document = Document(
            doc_id=doc_id,
            doc_type=doc_type,
            title=_title(doc_type, context.spec),
            agent_id=context.spec.agent_id,
            agent_version=context.spec.version,
            content=content,
            content_sha256=sha256_text(content),
            generated_at=context.generated_at,
            version=version,
            status=DocumentStatus.DRAFT,
            template=template_name,
            evidence_refs=sorted({record.evidence_id for record in context.evidence}),
            run_id=context.run.run_id if context.run else None,
        )

        if self._vault is not None:
            self._vault.put_text(
                "document",
                content,
                content_type="text/markdown",
                agent_id=document.agent_id,
                run_id=document.run_id,
                metadata={"doc_id": doc_id, "doc_type": doc_type.value},
            )
        if self._audit is not None:
            self._audit.append(
                actor="system",
                action="document.generated",
                entity_type="document",
                entity_id=doc_id,
                payload={
                    "doc_type": doc_type.value,
                    "content_sha256": document.content_sha256,
                    "template": template_name,
                },
            )
        return document

    def generate_package(
        self,
        spec: AgentSpec,
        *,
        doc_types: Sequence[DocumentType] | None = None,
        **components: Any,
    ) -> list[Document]:
        """Render the whole package in reading order.

        A document type whose inputs are absent is skipped rather than failing
        the package: a run that has not happened yet cannot produce an OQ
        report, and refusing to generate the URS because of it would help
        nobody. What is skipped is reported by :meth:`skipped`.
        """
        self._skipped: dict[str, str] = {}
        documents: list[Document] = []
        for doc_type in doc_types or PACKAGE_ORDER:
            try:
                context = build_context(spec, doc_type, clock=self._clock, **components)
            except DocumentError as error:
                self._skipped[doc_type.value] = str(error)
                continue
            documents.append(self.generate(doc_type, context))
        return documents

    def skipped(self) -> dict[str, str]:
        """Document types omitted from the last package, and why."""
        return dict(getattr(self, "_skipped", {}))

    # -- rendering ---------------------------------------------------------

    @staticmethod
    def render_signature_block(document: Document, block: str) -> Document:
        """Append the 11.50(b) human-readable signature manifest to a document.

        Returns a new document whose digest covers the signature block, so what
        was signed and what is displayed are the same bytes.
        """
        content = document.content.rstrip("\n") + "\n\n" + block.rstrip("\n") + "\n"
        return document.replace(content=content, content_sha256=sha256_text(content))

    def to_html(self, document: Document, *, standalone: bool = True) -> str:
        """Render a document as self-contained, printable HTML."""
        body = markdown_to_html(document.content)
        if not standalone:
            return body
        return _HTML_SHELL.format(
            title=html_module.escape(f"{document.doc_id} — {document.title}"),
            style=_PRINT_STYLE,
            body=body,
        )

    def to_docx(self, document: Document, path: str | Path) -> Path:
        """Write a document as .docx, if the optional extra is installed.

        The PDF/A route a regulated deployment usually wants is this file
        converted by LibreOffice in headless mode
        (``soffice --convert-to pdf:writer_pdf_Export``); that conversion is
        deliberately left outside the tool, since it is a deployment
        dependency rather than a library one.
        """
        try:
            import docx  # noqa: PLC0415 - deliberately lazy
        except ImportError as error:
            raise DocumentError(
                "writing .docx requires python-docx, which is not installed. "
                "Install it with: pip install 'valkit[docx]'"
            ) from error

        out = Path(path)
        builder = docx.Document()
        for line in document.content.split("\n"):
            stripped = line.rstrip()
            if stripped.startswith("#"):
                level = len(stripped) - len(stripped.lstrip("#"))
                builder.add_heading(stripped[level:].strip(), level=min(level, 4))
            elif stripped.startswith("- "):
                builder.add_paragraph(stripped[2:], style="List Bullet")
            elif stripped:
                builder.add_paragraph(stripped)
        builder.save(str(out))
        return out


def _title(doc_type: DocumentType, spec: AgentSpec) -> str:
    names = {
        DocumentType.URS: "User Requirements Specification",
        DocumentType.FRS: "Functional Requirements Specification",
        DocumentType.RISK_ASSESSMENT: "Risk Assessment",
        DocumentType.VALIDATION_PLAN: "Validation Plan",
        DocumentType.CREDIBILITY_PLAN: "Credibility Assessment Plan",
        DocumentType.CREDIBILITY_REPORT: "Credibility Assessment Report",
        DocumentType.IQ_PROTOCOL: "Installation Qualification Protocol",
        DocumentType.IQ_REPORT: "Installation Qualification Report",
        DocumentType.OQ_PROTOCOL: "Operational Qualification Protocol",
        DocumentType.OQ_REPORT: "Operational Qualification Report",
        DocumentType.PQ_PROTOCOL: "Performance Qualification Protocol",
        DocumentType.PQ_REPORT: "Performance Qualification Report",
        DocumentType.RTM: "Requirements Traceability Matrix",
        DocumentType.VSR: "Validation Summary Report",
        DocumentType.PERIODIC_REVIEW: "Periodic Review",
        DocumentType.CHANGE_CONTROL: "Change Control Record",
        DocumentType.TOOL_QUALIFICATION: "Tool Qualification",
    }
    return f"{names[doc_type]} — {spec.agent_id} v{spec.version}"


def _normalise(content: str) -> str:
    """Collapse runs of blank lines and guarantee a single trailing newline.

    Templates with conditional blocks leave ragged whitespace that varies with
    which branches ran. Normalising here keeps the digest stable against
    changes that are invisible on the page.
    """
    content = re.sub(r"\n{3,}", "\n\n", content.replace("\r\n", "\n"))
    return content.rstrip("\n") + "\n"


# --------------------------------------------------------------------------
# Markdown to HTML
# --------------------------------------------------------------------------

_INLINE_CODE = re.compile(r"`([^`]+)`")
_BOLD = re.compile(r"\*\*([^*]+)\*\*")
_ITALIC = re.compile(r"(?<![*\w])\*([^*\n]+)\*(?!\*)")
_LINK = re.compile(r"\[([^\]]+)\]\(([^)]+)\)")


def _inline(text: str) -> str:
    """Apply inline Markdown to an already-escaped fragment."""
    text = _INLINE_CODE.sub(lambda m: f"<code>{m.group(1)}</code>", text)
    text = _BOLD.sub(lambda m: f"<strong>{m.group(1)}</strong>", text)
    text = _ITALIC.sub(lambda m: f"<em>{m.group(1)}</em>", text)
    text = _LINK.sub(lambda m: f'<a href="{m.group(2)}">{m.group(1)}</a>', text)
    return text


def markdown_to_html(markdown: str) -> str:
    """Convert the Markdown subset the templates emit into HTML.

    Deliberately small: headings, paragraphs, lists, tables, fenced code,
    blockquotes, horizontal rules and inline emphasis. A general Markdown
    library would be another supplier to assess for a conversion the templates
    fully control, and an unsupported construct here is a template to fix
    rather than a dependency to add.
    """
    lines = markdown.replace("\r\n", "\n").split("\n")
    out: list[str] = []
    index = 0
    in_list = False
    in_code = False

    def close_list() -> None:
        nonlocal in_list
        if in_list:
            out.append("</ul>")
            in_list = False

    while index < len(lines):
        line = lines[index]
        stripped = line.strip()

        if stripped.startswith("```"):
            close_list()
            if in_code:
                out.append("</code></pre>")
                in_code = False
            else:
                out.append("<pre><code>")
                in_code = True
            index += 1
            continue

        if in_code:
            out.append(html_module.escape(line))
            index += 1
            continue

        if not stripped:
            close_list()
            index += 1
            continue

        if re.fullmatch(r"-{3,}|\*{3,}|_{3,}", stripped):
            close_list()
            out.append("<hr>")
            index += 1
            continue

        heading = re.match(r"(#{1,6})\s+(.*)", stripped)
        if heading:
            close_list()
            level = len(heading.group(1))
            out.append(f"<h{level}>{_inline(html_module.escape(heading.group(2)))}</h{level}>")
            index += 1
            continue

        # A table: a header row, a separator row, then body rows.
        if stripped.startswith("|") and index + 1 < len(lines) and re.fullmatch(
            r"\|[\s:\-|]+\|", lines[index + 1].strip()
        ):
            close_list()
            header = _split_row(stripped)
            out.append("<table><thead><tr>")
            out.extend(f"<th>{_inline(html_module.escape(cell))}</th>" for cell in header)
            out.append("</tr></thead><tbody>")
            index += 2
            while index < len(lines) and lines[index].strip().startswith("|"):
                out.append("<tr>")
                out.extend(
                    f"<td>{_inline(html_module.escape(cell))}</td>"
                    for cell in _split_row(lines[index].strip())
                )
                out.append("</tr>")
                index += 1
            out.append("</tbody></table>")
            continue

        if stripped.startswith("> "):
            close_list()
            out.append(f"<blockquote>{_inline(html_module.escape(stripped[2:]))}</blockquote>")
            index += 1
            continue

        bullet = re.match(r"[-*]\s+(.*)", stripped)
        if bullet:
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{_inline(html_module.escape(bullet.group(1)))}</li>")
            index += 1
            continue

        numbered = re.match(r"\d+\.\s+(.*)", stripped)
        if numbered:
            if not in_list:
                out.append("<ul>")
                in_list = True
            out.append(f"<li>{_inline(html_module.escape(numbered.group(1)))}</li>")
            index += 1
            continue

        close_list()
        out.append(f"<p>{_inline(html_module.escape(stripped))}</p>")
        index += 1

    close_list()
    if in_code:
        out.append("</code></pre>")
    return "\n".join(out)


def _split_row(line: str) -> list[str]:
    cells = line.strip().strip("|").split("|")
    return [cell.strip().replace("\\|", "|") for cell in cells]


_PRINT_STYLE = """
:root { color-scheme: light; }
body { font-family: Georgia, 'Times New Roman', serif; line-height: 1.55;
       max-width: 46rem; margin: 2rem auto; padding: 0 1.25rem; color: #16181d;
       background: #fff; }
h1, h2, h3, h4 { font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI',
       Roboto, Helvetica, Arial, sans-serif; line-height: 1.25; color: #0d0f13; }
h1 { font-size: 1.7rem; border-bottom: 2px solid #16181d; padding-bottom: .4rem; }
h2 { font-size: 1.3rem; margin-top: 2.2rem; border-bottom: 1px solid #d6d9e0;
     padding-bottom: .25rem; }
h3 { font-size: 1.08rem; margin-top: 1.6rem; }
table { border-collapse: collapse; width: 100%; margin: 1rem 0; font-size: .9rem;
        font-family: -apple-system, BlinkMacSystemFont, 'Segoe UI', sans-serif; }
th, td { border: 1px solid #ccd0d8; padding: .4rem .55rem; text-align: left;
         vertical-align: top; }
th { background: #f2f4f7; font-weight: 600; }
code { font-family: ui-monospace, SFMono-Regular, Menlo, monospace; font-size: .86em;
       background: #f2f4f7; padding: .1em .3em; border-radius: 3px; }
pre { background: #f7f8fa; border: 1px solid #e2e5ea; padding: .8rem; overflow-x: auto; }
pre code { background: none; padding: 0; }
blockquote { border-left: 3px solid #ccd0d8; margin-left: 0; padding-left: 1rem;
             color: #4a4f58; }
hr { border: none; border-top: 1px solid #d6d9e0; margin: 2rem 0; }
@media print {
  body { max-width: none; margin: 0; font-size: 10.5pt; }
  h1, h2, h3 { page-break-after: avoid; }
  table, pre, blockquote { page-break-inside: avoid; }
}
"""

_HTML_SHELL = """<!doctype html>
<html lang="en">
<head>
<meta charset="utf-8">
<meta name="viewport" content="width=device-width, initial-scale=1">
<title>{title}</title>
<style>{style}</style>
</head>
<body>
{body}
</body>
</html>
"""
